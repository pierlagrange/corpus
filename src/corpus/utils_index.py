import io
import re
from typing import Any

# import camelot
from unicodedata import normalize

import json2html
import pandas as pd
import pylibmagic
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.table import Table
from langchain.text_splitter import RecursiveCharacterTextSplitter, SpacyTextSplitter
from magic import from_buffer
from pypdf import PdfReader

LIBMAGIC_VERSION = pylibmagic._version
# https://medium.com/@alice.yang_10652/convert-word-doc-or-docx-to-pdf-with-python-a-comprehensive-guide-6c8e8b5a079a#0ed0
# Ref mime/type : https://developer.mozilla.org/en-US/docs/Web/HTTP/Basics_of_HTTP/MIME_types/Common_types

WORD_EXTENSION = ["vnd.openxmlformats-officedocument.wordprocessingml.document", "CDFV2"]
PDF_EXTENSION = ["pdf"]
EXCEL_EXTENSION = ["vnd.ms-excel", "vnd.openxmlformats-officedocument.spreadsheetml.sheet"]
PPT_EXTENSION = ["vnd.ms-powerpoint"]
CHUNK_OPTIONS = ["tiktoken_recursive", "recursive", "spacy"]
WEB_EXTENSION = ["html"]
WEB_REGEXP = r"https?://(?:www\.)?[a-zA-Z0-9./]+"


# File type utils
def split_format(s: str, filter_source: bool = True) -> str | tuple[str, str] | Exception:
    """
    A function that splits a string into source and format based on the '/' delimiter.

    Parameters:
        s (str): The input string to be split.
        filter_source (bool, optional): A flag to determine whether to return only the format or both source and format. Defaults to True.

    Returns:
        Tuple[str, str] or Exception: If filter_source is True, returns the format.
                                     If filter_source is False, returns a tuple of source and format.
                                     If an exception occurs during the split operation, returns the exception object.
    """

    try:
        source, file_format = s.split("/")
    except Exception as e:
        raise e
    if filter_source:
        return file_format
    else:
        return source, file_format


def is_valid_url(url):
    """
    A function to check if a given URL is valid based on a regular expression pattern.

    Parameters:
    url (str): The URL to be validated.

    Returns:
    bool: True if the URL is valid, False otherwise.
    """
    url_regex = re.compile(WEB_REGEXP)
    return bool(url_regex.match(url))


# Files content utils


def word_to_str(raw: bytes, file_type: str) -> dict | ValueError:
    """
    Convert raw bytes to a string representation based on the file type.

    Args:
        raw (bytes): The raw bytes to be converted.
        type (str): The type of the file.

    Returns:
        Dict: A dictionary where the key is the paragraph id and the value is the text content of that paragraph.
    """
    if file_type in WORD_EXTENSION:
        res = {}
        doc = Document(raw)
        for ids, paragraph in enumerate(doc.paragraphs):
            res[ids] = paragraph.text
        return res
    else:
        raise ValueError(f"Invalid type : {file_type}")


def html_to_str(url: str) -> dict | ValueError:
    """
    Fetches the HTML content of a given URL and converts it to a string.
    Args:
        url (str): The URL to fetch the HTML content from.
    Returns:
        dict: A dictionary where the key is the url and the value is the HTML content as a string.
    Raises:
        ValueError: If the URL is not valid.
        Exception: If there's any other error when fetching the HTML content.
    """
    try:
        if is_valid_url(url):
            r = requests.get(url, timeout=10)
            soup = BeautifulSoup(r.content)
            res = {url: soup.get_text(separator="\n", strip=True)}
            return res
        else:
            raise ValueError(f"{url} is not a valid url")
    except Exception as e:
        raise e


def pdf_to_str(raw: bytes, file_type: str) -> dict | ValueError:
    """
    Convert raw bytes to a string representation based on the file type.

    Args:
        raw (bytes): The raw bytes to be converted.
        type (str): The type of the file.

    Returns:
        Dict: A dictionary where the keys are page numbers and the values are the extracted text from each page.

    Raises:
        ValueError: If the file type is not supported.

    """

    if file_type in PDF_EXTENSION:
        res = {}
        if file_type in PDF_EXTENSION:
            doc = PdfReader(raw)
            for num_page, page in enumerate(doc.pages):
                res[num_page] = page.extract_text()
        return res
    else:
        raise ValueError(f"Invalid type : {file_type}")


def excel_to_table(raw: bytes, file_type: str) -> dict:
    """
    Convert raw Excel bytes into a dictionary of dataframes, where each key is a sheet name and the value is a dataframe of the sheet's data.

    Args:
        raw (bytes): The raw Excel bytes to be converted.
        type (str): The type of the file.

    Returns:
        Dict: A dictionary where the keys are sheet names and the values are dataframes of the sheet's data.

    Raises:
        ValueError: If the file type is not supported.

    """

    if file_type in EXCEL_EXTENSION:
        res: dict = {}
        reader = pd.ExcelFile(raw)
        for sheet in reader.sheet_names:
            res = {**res, **{sheet: reader.parse(sheet, keep_default_na=False).to_dict(orient="records")}}
        return res
    else:
        raise ValueError(f"Invalid type : {file_type}")


# Chunks utils
def spacy_chunk(s: str) -> list | Exception:
    """
    Splits a given string into smaller chunks based on a spacy linguistic model (fr_core_news_sm).

    Args:
        s (str): The string to be split.
        chunk_size (int, optional): The size of each chunk. Defaults to 10000.

    Returns:
        list: A list of smaller chunks of the given string.
    """
    try:
        splitter = SpacyTextSplitter(pipeline="fr_core_news_sm", separator="\n\n")
        splits = splitter.split_text(s)
        return splits
    except Exception as e:
        raise e from e


def recursive_chunk(s: str, recursive_coeff: float = 0.1, chunk_size: int = 800) -> list | Exception:
    """
    Splits a given string into smaller chunks based on a recursive character text splitter.

    Args:
        s (str): The string to be split.
        recursive_coeff (float, optional): The coefficient for the recursive split. Defaults to 0.1.
        chunk_size (int, optional): The size of each chunk. Defaults to 800.

    Returns:
        list: A list of smaller chunks of the given string.
    """
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=recursive_coeff, separators=["\n\n"]
        )
        splits = splitter.split_text(s)
        return splits
    except Exception as e:
        raise e from e


def tiktoken_recursive_chunk(s: str, recursive_coeff: float = 0.1, chunk_size: int = 800) -> list:
    """
    Splits a given string into smaller chunks using a TikTok encoder-based RecursiveCharacterTextSplitter.

    Args:
        s (str): The string to be split.
        recursive_coeff (float, optional): The coefficient for the recursive split. Defaults to 0.1.
        chunk_size (int, optional): The size of each chunk. Defaults to 200.

    Returns:
        list: A list of smaller chunks of the given string.
    """
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        model_name="gpt-4", chunk_size=chunk_size, chunk_overlap=recursive_coeff, separators=["\n\n"]
    )
    splits = splitter.split_text(s)
    return splits


# Gestion des tables
# https://www.microsoft.com/en-us/research/blog/improving-llm-understanding-of-structured-data-and-exploring-advanced-prompting-methods/?msockid=23c103485e6b6fa5237317c05f2c6e06
# https://www.microsoft.com/en-us/research/uploads/prod/2023/12/wsdm24-SUC.pdf


def describe_table(key) -> str:
    """Not implemented"""
    return key


def process_table(table: dict) -> bytes | str:
    """
    Convert a dictionary representing a table to an HTML table string.

    Args:
        table (Dict): A dictionary representing a table. Each key-value pair in the dictionary represents a row in the table, where the key is the row index and the value is a list of cell values for that row.

    Returns:
        str: An HTML table string representing the input table.
    """

    html_table = json2html.json2html.convert(table)
    return html_table


def enumerate_word_table(table: Table) -> dict:
    """
    Convert a Table object into a dictionary where each key is a row index and the corresponding value is a list of cell values for that row.

    Args:
        table (Table): The input Table object to be converted.

    Returns:
        Dict: A dictionary representation of the input table where each key-value pair represents a row in the table.
    """

    num_row = range(len(table.rows))
    num_col = range(len(table.columns))
    res = {idr: "/t".join([normalize("NFKD", table.cell(idr, idc).text) for idc in num_col]) for idr in num_row}
    return res


def extract_table_word(raw: bytes, file_type: str) -> dict[Any, Any] | ValueError:
    """
    Extracts table data from a document in bytes format based on the file type.

    Args:
        raw (bytes): The raw document in bytes.
        file_type (str): The type of the document.

    Returns:
        Dict: A dictionary where the keys are table indices and the values are dictionaries representing the tables.

    Raises:
        ValueError: If the file type is not in the WORD_EXTENSION list.
    """

    if file_type in WORD_EXTENSION:
        res: dict = {}
        doc = Document(raw)
        for idt, table in enumerate(doc.tables):
            res[idt] = enumerate_word_table(table)
        return res
    else:
        return ValueError(f"Invalid type : {type}")


def extract_table_url(raw: BeautifulSoup, file_type: str) -> dict | ValueError:
    """
    Extracts table data from a document in bytes format based on the file type.

    Args:
        raw (bytes): The raw document in bytes.
        file_type (str): The type of the document.

    Returns:
        Dict: A dictionary where the keys are table indices and the values are dictionaries representing the tables.

    Raises:
        ValueError: If the file type is not in the WORD_EXTENSION list.
    """

    if file_type in WORD_EXTENSION:
        res: dict = {}
        doc = Document(raw)
        for idt, table in enumerate(doc.tables):
            res[idt] = enumerate_word_table(table)
        return res
    else:
        return ValueError(f"Invalid type : {type}")


# def extract_table_pdf(raw_pdf:bytes)->Dict:
#     """
#     Extracts tables from a PDF file and returns a dictionary containing the extracted tables.

#     Args:
#         raw_pdf (bytes): The raw PDF file.

#     Returns:
#         dict: A dictionary containing the extracted tables. The keys are the indices of the tables, and the values are lists of strings representing the rows of the tables. Each row is a string where the values are separated by '/t'.

#     """

#     extract = camelot.read_pdf(raw_pdf)
#     page_tables = [table for table in extract]
#     res = {}
#     if len(page_tables)>0:
#         tables = [t.df.to_dict(orient = 'records') for t in page_tables]
#         for idx,table in enumerate(tables):
#             table_res = {}
#             table_res[idx] = []
#             for line in table:
#                 table_res[idx].append("/t".join([normalize('NFKD',x) for x in line.values()]))
#             res = {**res,
#                    **table_res}
#     return res


# Classes
class Chunker:
    """
    The Chunker class is responsible for breaking down a given string into smaller chunks. The chunking strategy
    can be specified during the initialization of the class.

    Attributes:
        option (str): The chunking strategy to use. It can be either 'recursive', 'tiktoken_recursive', or 'spacy'.
        chunks (list[str]): A list to store the chunks of the string.

    The class provides a method, 'create_chunks', to break down a given string into chunks based on the specified
    chunking strategy. If an unsupported chunking strategy is specified, a 'NotImplementedError' is raised.

    Raises:
        ValueError: If the provided chunking option is not in the CHUNK_OPTIONS list.
        NotImplementedError: If the specified chunking option is not supported.
    """

    def __init__(self, option: str = "recursive") -> None:
        """
        Initializes a new instance of the Chunker class.

        Args:
            option (str, optional): The chunking option to use. Defaults to "recursive".
            It can be either 'recursive', 'tiktoken_recursive', or 'spacy'

        Raises:
            ValueError: If the provided option is not in the CHUNK_OPTIONS list.

        Returns:
            None
        """
        if option in CHUNK_OPTIONS:
            self.option = option
        else:
            raise ValueError(f"Option {option} not in {CHUNK_OPTIONS}")
        self.chunks: list[str] = []

    def create_chunks(self, s):
        """
        Creates chunks of a given string based on the specified option.

        Args:
            s (str): The string to be chunked.

        Returns:
            list: A list of chunks of the given string.

        Raises:
            NotImplementedError: If the specified option is not supported.
        """

        if self.option == "recursive":
            self.chunks = recursive_chunk(s)
        elif self.option == "tiktoken_recursive":
            self.chunks = tiktoken_recursive_chunk(s)
        elif self.option == "spacy":
            self.chunks = spacy_chunk(s)
        else:
            raise NotImplementedError
        return self.chunks


class Ingestor:
    """
    The Ingestor class is responsible for handling raw data, guessing its format,
    processing the documents and creating chunks of text from the documents.

    It provides methods to guess the format of raw data, process the documents based
    on their file type, and create chunks of text from the processed documents.

    Attributes:
        raw (dict): A dictionary containing the raw data.
        chunker (Chunker): An instance of the Chunker class.
        extract_table_option (bool): A flag to indicate whether to extract tables from the document.
        format (dict): A dictionary to store the format of the raw data.
        output (dict): A dictionary to store the processed data.
        chunks_output (dict): A dictionary to store the chunks of text from the documents.
        failed_to_process (dict): A dictionary to store any exceptions that occurred during processing.
    """

    def __init__(self, raw: dict[str, bytes], chunker: Chunker, extract_table: bool = False) -> None:
        """
        Initializes a new instance of the class.
        Args:
            raw (Dict[str, bytes]): A dictionary of raw data.
            chunker (Chunker): An instance of the Chunker class.
            extract_table (bool, optional): A flag to indicate whether to extract tables from the document. Defaults to False.
        Returns:
            None
        """
        self.failed_to_process: dict = {}
        self.raw: dict[str, Any] = raw
        for k, v in raw.items():
            if not isinstance(v, io.BufferedReader):
                try:
                    self.raw[k] = io.BytesIO(v)
                except Exception as e:
                    self.failed_to_process = {
                        **self.failed_to_process,
                        **{k: f"Fail to transform in BufferedReader {e}"},
                    }
        self.format: dict = {}
        self.output: dict = {}
        self.chunks_output: dict = {}
        self.extract_table_option = extract_table
        self.chunker = chunker

    def guess_format(self) -> None:
        """
        Guesses the format of the raw data.

        This function iterates over the items in the `raw` dictionary and tries to guess the format of each value. It does this by reading the first 5000 bytes of the value using the `from_buffer` function from the `magic` library. The `mime` parameter is set to `True` to get the MIME type of the data. The MIME type is then used to split the format using the `split_format` function.

        If an exception occurs during the process, the exception is caught and stored in the `self.format` dictionary under the key `k`.

        Parameters:
            None

        Returns:
            None
        """
        for k, v in self.raw.items():
            try:
                file_type = from_buffer(v.read(5000), mime=True)
                self.format = {**self.format, k: {"mime": file_type, "file_type": split_format(file_type)}}
            except Exception as e:
                self.failed_to_process = {**self.failed_to_process, **{k: e}}

    def process_docs(self, extract_table: bool = False) -> None:
        """
        Process the documents in the `self.raw` dictionary based on their file type.

        This function iterates over the items in the `self.raw` dictionary and checks the file type of each value.

        If the file type is in the `PDF_EXTENSION` list, it extracts tables from the PDF using the `extract_table_pdf` function and processes each table using the `process_table` function. The resulting table data is stored in the `dict_tables` dictionary with a key generated by the `describe_table` function. The `self.output` dictionary is updated with the processed data, including the full text of the PDF and the table data.

        If the file type is in the `EXCEL_EXTENSION` list, it processes the Excel file using the `excel_to_table` function. The resulting table data is stored in the `dict_tables` dictionary. The `self.output` dictionary is updated with the processed data, including the full text of the Excel file and the table data.

        Parameters:
            None

        Returns:
            None
        """
        for k, v in self.raw.items():
            try:
                file_type = self.format[k]["file_type"]
                if file_type in PDF_EXTENSION:
                    dict_tables: dict[Any, Any] = {}
                    # if extract_table:
                    # tables = extract_table_pdf(v)
                    # for key,table in tables.items():
                    #     dict_tables[describe_table(key)] = process_table(table)
                    self.output = {**self.output, **{k: {"full_text": pdf_to_str(v, file_type), "table": dict_tables}}}
                elif file_type in EXCEL_EXTENSION:
                    dict_tables = {}
                    tables = excel_to_table(v, file_type)
                    for key, table in tables.items():
                        dict_tables[describe_table(key)] = process_table(table)
                    self.output = {**self.output, **{k: {"full_text": None, "table": dict_tables}}}
                elif file_type in WORD_EXTENSION:
                    dict_tables = {}
                    if extract_table:
                        tables = extract_table_word(v, file_type)
                        for key, table in tables.items():
                            dict_tables[describe_table(key)] = process_table(table)
                    self.output = {**self.output, **{k: {"full_text": word_to_str(v, file_type), "table": dict_tables}}}
                else:
                    self.failed_to_process = {**self.failed_to_process, **{k: f"Format non-reconnu {file_type}"}}
            except Exception as e:
                self.failed_to_process = {**self.failed_to_process, **{k: e}}

    def create_chunks(self) -> None:
        """
        Create chunks of text from the output dictionary.

        This function iterates over the `self.output` dictionary and checks if each document has a 'full_text' key.
        If it does, it creates chunks of text using the `self.chunker.create_chunks()` method.
        The resulting chunks are stored in the `chunked_text` dictionary.

        The function then updates the `self.chunks_output` dictionary with the new chunks and the corresponding 'table' key.

        Parameters:
            self (object): The current instance of the class.

        Returns:
            None
        """

        if len(self.output) == 0:
            self.process_docs()
        try:
            for id_doc, doc in self.output.items():
                if doc["full_text"]:
                    chunked_text = {}
                    for page, content in doc["full_text"].items():
                        chunked_text[page] = self.chunker.create_chunks(content)
                    self.chunks_output = {
                        **self.chunks_output,
                        **{id_doc: {"full_text": chunked_text, "table": doc["table"]}},
                    }
                else:
                    self.chunks_output = {**self.chunks_output, **{id_doc: {"full_text": None, "table": doc["table"]}}}
        except Exception as e:
            self.failed_to_process = {**self.failed_to_process, **{id_doc: e}}


class WebIngestor:
    """
    The WebIngestor class is responsible for handling web links, validating their format,
    processing the HTML content of the web links and creating chunks of text from the content.

    It provides methods to validate the format of web links, process the HTML content
    based on their file type, and create chunks of text from the processed content.

    Attributes:
        raw (dict): A dictionary containing the web links.
        chunker (Chunker): An instance of the Chunker class.
        extract_table_option (bool): A flag to indicate whether to extract tables from the website.
        format (dict): A dictionary to store the format of the web links.
        output (dict): A dictionary to store the processed data.
        chunks_output (dict): A dictionary to store the chunks of text from the content.
        failed_to_process (dict): A dictionary to store any exceptions that occurred during processing.
    """

    def __init__(self, raw: dict[str, str], chunker: Chunker, extract_table: bool = False) -> None:
        """
        Initializes a new instance of the class.

        Args:
            raw (Dict[str, str]): A dictionary of web link as string.
            chunker (Chunker): An instance of the Chunker class.

        Returns:
            None
        """
        self.raw = raw
        self.format: dict = {}
        self.output: dict = {}
        self.extract_table_option = extract_table
        self.chunker = chunker
        self.chunks_output: dict = {}
        self.failed_to_process: dict = {}

    def guess_format(self) -> None:
        """
        Guesses the format of the raw data.

        This function iterates over the items in the `raw` dictionary and tries to guess the format of each value. It does this by reading the first 5000 bytes of the value using the `from_buffer` function from the `magic` library. The `mime` parameter is set to `True` to get the MIME type of the data. The MIME type is then used to split the format using the `split_format` function.

        If an exception occurs during the process, the exception is caught and stored in the `self.format` dictionary under the key `k`.

        Parameters:
            None

        Returns:
            None
        """

        for k, v in self.raw.items():
            try:
                webtype = is_valid_url(v)
                if not webtype:
                    webe = "Please pass a valid http url to open"
                    self.failed_to_process = {**self.failed_to_process, **{k: webe}}
                    raise ValueError(webe)
                else:
                    self.format = {**self.format, k: {"is_valid": webtype}}
            except Exception as e:
                self.failed_to_process = {**self.failed_to_process, **{k: e}}

    def process_docs(self, extract_table: bool = False) -> None:
        """
        Process the documents in the `self.raw` dictionary based on their file type.

        This function iterates over the items in the `self.raw` dictionary and checks the file type of each value. If the file type is in the `WORD_EXTENSION` list, it extracts tables from the Word document using the `extract_table_word` function and processes each table using the `process_table` function. The resulting table data is stored in the `dict_tables` dictionary with a key generated by the `describe_table` function. The `self.output` dictionary is updated with the processed data, including the full text of the Word document and the table data.

        Parameters:
            extract_table (bool, optional): A flag to indicate whether to extract tables from the Word document. Defaults to False.

        Returns:
            None
        """

        try:
            for k, v in self.raw.items():
                valid_url = self.format[k]["is_valid"]
                if valid_url:
                    dict_tables: dict = {}
                    if extract_table:
                        pass
                        # tables = extract_table_word(v,file_type)
                        # for key,table in tables.items():
                        #     dict_tables[describe_table(key)] = process_table(table)
                    self.output = {**self.output, **{k: {"full_text": html_to_str(v), "table": dict_tables}}}
        except Exception as e:
            self.failed_to_process = {**self.failed_to_process, **{k: e}}

    def create_chunks(self) -> None:
        """
        Create chunks of text from the output dictionary.

        This function iterates over the `self.output` dictionary and checks if each document has a 'full_text' key.
        If it does, it creates chunks of text using the `self.chunker.create_chunks()` method.
        The resulting chunks are stored in the `chunked_text` dictionary.

        The function then updates the `self.chunks_output` dictionary with the new chunks and the corresponding 'table' key.

        Parameters:
            self (object): The current instance of the class.

        Returns:
            None
        """

        if len(self.output) == 0:
            self.process_docs()

        for id_doc, doc in self.output.items():
            try:
                if doc["full_text"]:
                    chunked_text = {}
                    for page, content in doc["full_text"].items():
                        chunked_text[page] = self.chunker.create_chunks(content)
                    self.chunks_output = {
                        **self.chunks_output,
                        **{id_doc: {"full_text": chunked_text, "table": doc["table"]}},
                    }
                else:
                    self.chunks_output = {**self.chunks_output, **{id_doc: {"full_text": None, "table": doc["table"]}}}
            except Exception as e:
                self.failed_to_process = {**self.failed_to_process, **{id_doc: e}}
